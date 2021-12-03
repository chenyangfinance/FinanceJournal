# coding:utf-8
import urllib.request
import re
import deepl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def intToR(num):  # 数字转罗马数字,copy自 https://www.jianshu.com/p/c89142aa1cb0
    c = {
        'g': ('', 'I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX'),
        's': ('', 'X', 'XX', 'XXX', 'XL', 'L', 'LX', 'LXX', 'LXXX', 'XC'),
        'b': ('', 'C', 'CC', 'CCC', 'CD', 'D', 'DC', 'DCC', 'DCCC', 'CM'),
        'q': ('', 'M', 'MM', 'MMM')
    }
    roman = []
    roman.append(c['q'][num // 1000])
    roman.append(c['b'][(num // 100) % 10])
    roman.append(c['s'][(num // 10) % 10])
    roman.append(c['g'][num % 10])
    return ''.join(roman)


# 初始化deepL
translator = deepl.Translator("e3c66233-3860-6d4e-22dc-eabe4408a3ca:fx")

url = 'https://www.sciencedirect.com/journal/journal-of-banking-and-finance/vol/133/suppl/C'
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/72.0.3626.119 Safari/537.36',
}
request = urllib.request.Request(url=url, headers=headers)
content = urllib.request.urlopen(request).read().decode('utf8')
#print(content)

# 0. 文档信息
#<title data-react-helmet="true">Journal of Banking &amp; Finance | Vol 133, December 2021 | ScienceDirect.com by Elsevier</title>
pattern0 = re.compile(r'<title data-react-helmet="true">(.*?)ScienceDirect.com')
date = pattern0.findall(content)[0].split("| ")[1]
issue = date.split(", ")[0] + " " + date.split(", ")[1]

# 1. title list
# <span class="js-article-title">Integrating corporate social responsibility criteria into executive compensation and firm innovation: International evidence</span></span></a></h3></dt><dd
pattern1 = re.compile(r'<span class="js-article-title">(.*?)</span></span>')
titleList = pattern1.findall(content)[1:]  # 删除一个Editorial Board
print(len(titleList))
# print(titleList)

# 2. author list
# <div class="text-s u-clr-grey8 js-article__item__authors">Albert Tsang, Kun Tracy Wang, Simeng Liu, Li Yu</div></dd>
pattern2 = re.compile(r'<div class="text-s u-clr-grey8 js-article__item__authors">(.*?)</div>')
authorList = pattern2.findall(content)[0:]
print(len(authorList))
# print(authorList)

# 3. id list 来生成每篇文章的URL,之后访问URL读取Abstract
# for="checkbox-S0165410121000434"><input type="checkbox
pattern3 = re.compile(r'for="checkbox-(.*?)"><input type="checkbox')
idList = pattern3.findall(content)[1:]  # 删除第一个Editorial Board
print(len(idList))
print(idList)
# 4. 循环idList 获取摘要
# Abstract</h2><div id="as0005"><p id="sp0115">I use the staggered adoption of state-level antitakeover laws to provide causal evidence that managerial agency problems reduce the allocative efficiency of conglomerate firms. I find that increases in control slack following the passage of antitakeover laws reduces <em>q</em>-sensitivity of investment by 64%. The adverse impact of the laws appears mostly at conglomerate firms that benefited from disciplinary takeover threats prior to the passage of the laws, lacked alternative sources of pressure on management, or had the structural makings to fuel wasteful influence activities and power struggles among managers. These findings suggest that takeover threats impact the efficiency of resource allocation.</p></div></div></div>
absList = []
abs_url = r"https://www.sciencedirect.com/science/article/pii/"
pattern4 = re.compile(r'A\D\D\D\D\D\D\D</h2><div id="\D+\d+"><p id="\D+\d+">(.*?)</p></div></div></div>')

for i in range(0, len(idList)):  # len(idList)
    try:
        print("正在获取摘要(%d)..." % (i + 1))
        id = abs_url + idList[i]
        print(id)
        # https://www.sciencedirect.com/science/article/pii/S0929119921001929
        abs_content = urllib.request.Request(url=id, headers=headers)
        abs = urllib.request.urlopen(abs_content).read().decode('utf8')
        # print(abs)
        a = pattern4.findall(abs)[0]  # 和JAE的不同，这里不切分字符串，而是用正则提取
        a = a.replace("&#x27;", "'")  # 单引号乱码
        a = re.sub(r'<.*?>', '', a)  # html乱码
        print(a)
        absList.append(a)
    except:
        a = " "
        print(a)
        absList.append(a)
        print("正在获取摘要(%d)..." % (i + 1))


# print(absList)

# 5. 写入Word
document = Document()
document.add_paragraph("刊名: Journal of Banking and Finance")
document.add_paragraph("刊号: " + date)
document.add_paragraph("仅翻译用于学术交流，版权归期刊和作者所有")
document.add_paragraph("原网页: " + url)
document.add_paragraph(" ")
document.add_paragraph(" ")
document.add_paragraph(" ")

for i in range(0, len(idList)):
    print(i)
    p1 = document.add_paragraph()
    run1 = p1.add_run(str(i + 1) + ". " + titleList[i])  # 使用add_run添加文字
    p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
    run1.bold = True  # 字体加粗

    p5 = document.add_paragraph()
    cn_title = translator.translate_text(titleList[i], target_lang="ZH")
    run5 = p5.add_run(str(cn_title))
    p5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
    document.add_paragraph(" ")

    p2 = document.add_paragraph()
    run2 = p2.add_run("作者: " + authorList[i])  # 使用add_run添加文字
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 段落文字左对齐
    document.add_paragraph(" ")

    p3 = document.add_paragraph()
    run3 = p3.add_run("摘要: " + absList[i])  # 使用add_run添加文字
    p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 段落文字平铺设置

    p6 = document.add_paragraph()
    run6 = p6.add_run()  # 使用add_run添加文字
    cn_abs = translator.translate_text(absList[i], target_lang="ZH")
    run6 = p6.add_run(str(cn_abs))
    p6.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 段落文字居中设置

    p7 = document.add_paragraph()
    run7 = p7.add_run("原文链接: " + abs_url + idList[i])  # 使用add_run添加文字
    p7.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 段落文字居中设置
    run7.font.size = Pt(9)  # 字体大小设置，和word里面的字号相对应，小一
    run7.italic = True  # 斜体
    document.add_paragraph(" ")
    document.add_paragraph(" ")
    document.add_paragraph(" ")
    document.add_paragraph(" ")
    document.add_paragraph(" ")

p8 = document.add_paragraph()
run8 = p8.add_run(r"(关于本次推送：我编写代码从期刊官网摘取论文的题目，作者，和摘要，并调用 DeepL 的接口翻译题目和摘要。我仅对翻译内容进行了非常粗略的审阅和修正。代码公布在我的Github页面：https://github.com/chenyangfinance/FinanceJournal)" )  # 使用add_run添加文字
p8.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 段落文字居中设置
run8.font.size = Pt(9)  # 字体大小设置，和word里面的字号相对应，小一
run8.italic = True  # 斜体

document.save('JBF %s.docx' % issue)

