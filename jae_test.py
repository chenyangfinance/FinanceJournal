#coding:utf-8
import urllib.request
import re
import deepl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def intToR(num):  # 数字转罗马数字,copy自 https://www.jianshu.com/p/c89142aa1cb0
    c = {
        'g': ('', 'I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX'),
        's': ('', 'X', 'XX', 'XXX', 'XL', 'L', 'LX', 'LXX', 'LXXX', 'XC'),
        'b': ('', 'C', 'CC', 'CCC', 'CD', 'D', 'DC', 'DCC', 'DCCC', 'CM'),
        'q': ('', 'M', 'MM', 'MMM')
    }
    roman = []
    roman.append(c['q'][num // 1000])
    roman.append(c['b'][(num // 100) % 10])
    roman.append(c['s'][(num // 10) % 10])
    roman.append(c['g'][num % 10])
    return ''.join(roman)

#初始化deepL
translator = deepl.Translator("e3c66233-3860-6d4e-22dc-eabe4408a3ca:fx")

url = 'https://www.sciencedirect.com/journal/journal-of-accounting-and-economics/vol/72/issue/2'
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/72.0.3626.119 Safari/537.36',
}
request = urllib.request.Request(url=url, headers=headers)
content = urllib.request.urlopen(request).read().decode('utf8')
#print(content)

# 0. 文档信息
#<title data-react-helmet="true">JAE | Journal of Accounting and Economics | Vol 72, Issues 2–3, November–December 2021 | ScienceDirect.com
pattern0 = re.compile(r'<title data-react-helmet="true">JAE(.*?)ScienceDirect.com')
date = pattern0.findall(content)[0].split("| ")[2]
issue = date.split(", ")[0] + " " + date.split(", ")[1]

# 1. title list
# <span class="js-article-title">Cash-based bonus plans as a strategic communication, coordination and commitment mechanism</span></span></a></h3></dt>
pattern1 = re.compile(r'<span class="js-article-title">(.*?)</span></span>')
titleList = pattern1.findall(content)[2:-1] #删除前两个和最后一个，一个是Editorial Board, 一个是Editorial Data，最后有个Acknowledgement
#print(len(titleList))
#print(titleList)

# 2. author list
#<div class="text-s u-clr-grey8 js-article__item__authors">Matthew Bloomfield, Brandon Gipper, John D. Kepler, David Tsui</div>
pattern2 = re.compile(r'<div class="text-s u-clr-grey8 js-article__item__authors">(.*?)</div>')
authorList = pattern2.findall(content)
#print(len(authorList))
#print(authorList)

# 3. id list 来生成每篇文章的URL,之后访问URL读取Abstract
# for="checkbox-S0165410121000434"><input type="checkbox
pattern3 = re.compile(r'for="checkbox-(.*?)"><input type="checkbox')
idList = pattern3.findall(content)[2:-1] #删除前两个和最后一个，一个是Editorial Board, 一个是Editorial Data，最后有个Acknowledgement
print(len(idList))

# 4. 循环idList 获取摘要
# <p id="abspara0010">Executive bonus plans often incorporate performance measures that exclude particular costs—a practice we refer to as “cost shielding.” We predict that boards use cost shielding to mitigate underinvestment and insulate new managers from the costs of prior executives’ decisions. We find evidence that boards use cost shielding to deter underinvestment in intangibles and encourage managers to take advantage of growth opportunities. We also find that cost shielding tends to be elevated for newly-hired executives, and decreases over tenure. Collectively, our results suggest that boards deliberately choose performance metrics that alleviate agency conflicts.</p></div></div></div>
absList = []
abs_url = r"https://www.sciencedirect.com/science/article/pii/"
pattern4 = re.compile(r'<p id="abspara0010">(.*?)"></p></div></div></div>')

#https://www.sciencedirect.com/science/article/pii/S0165410121000434
for i in range(0, len(idList)):
    id = abs_url + idList[i]
    abs_content = urllib.request.Request(url=id, headers=headers)
    abs = urllib.request.urlopen(abs_content).read().decode('utf8')
    #print(type(abs))  # 字符串，所以按字符串处理，把abstract切出来
    a = abs.split('Abstract</h2><div id="abssec0010"><p id="abspara0010">')[-1]
    b = a.split('</p></div></div></div><ul id="issue-navigation"')[0] #the content of abstract
    absList.append(b)
    print("正在获取摘要(%d)..." % (i+1))
#print(absList)


# 5. 写入Word
document = Document()
document.add_paragraph("刊名: Journal of Accounting and Economics")
document.add_paragraph("刊号: " + date)
document.add_paragraph("仅翻译用于学术交流，版权归期刊和作者所有")
document.add_paragraph(" ")
document.add_paragraph(" ")
document.add_paragraph(" ")

for i in range(0, len(idList)):

    p1 = document.add_paragraph()
    run1 = p1.add_run(str(intToR(i+1))+". "+titleList[i])  # 使用add_run添加文字
    p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
    run1.bold = True  # 字体加粗
    '''
    p5 = document.add_paragraph()
    cn_title = translator.translate_text(titleList[i], target_lang="ZH")
    run5 = p5.add_run(str(cn_title))
    p5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
    document.add_paragraph(" ")
    '''
    p2 = document.add_paragraph()
    run2 = p2.add_run("作者: "+authorList[i])  # 使用add_run添加文字
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 段落文字居中设置    p1 = document.add_paragraph()
    document.add_paragraph(" ")

    p3 = document.add_paragraph()
    run3 = p3.add_run("摘要: "+absList[i])  # 使用add_run添加文字
    p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 段落文字居中设置
    '''
    p6 = document.add_paragraph()
    run6 = p6.add_run()  # 使用add_run添加文字
    cn_abs = translator.translate_text(absList[i], target_lang="ZH")
    run6 = p6.add_run(str(cn_abs))
    p6.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 段落文字居中设置
    '''
    p7 = document.add_paragraph()
    run7 = p7.add_run("原文链接: " + abs_url + idList[i])  # 使用add_run添加文字
    p7.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 段落文字居中设置
    run7.font.size = Pt(9)  # 字体大小设置，和word里面的字号相对应，小一
    run7.italic = True
    document.add_paragraph(" ")
    document.add_paragraph(" ")

document.save('JAE %s.docx' % issue)

